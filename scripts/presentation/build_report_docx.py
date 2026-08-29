"""Converts docs/FINAL_PROJECT_REPORT.md into a professionally formatted
.docx (headings, tables, code blocks, bullet lists) using a lightweight
custom parser tuned to this report's own markdown structure -- no pandoc/
libreoffice available in this environment.
"""
import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor, Inches
from docx.enum.table import WD_TABLE_ALIGNMENT

REPO = Path("/home/muhammad/Documents/smart-detector")
SRC = REPO / "docs" / "FINAL_PROJECT_REPORT.md"
OUT = Path(sys.argv[1]) if len(sys.argv) > 1 else REPO / "deliverables" / "Factory_Safety_Sentinel_Final_Report.docx"

NAVY = RGBColor(0x0A, 0x14, 0x20)
ORANGE = RGBColor(0xB4, 0x50, 0x09)
GRAY = RGBColor(0x47, 0x55, 0x69)
CODE_BG = "F1F5F9"

doc = Document()

# Base style
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)

for sec in doc.sections:
    sec.left_margin = Inches(1)
    sec.right_margin = Inches(1)


def set_cell_shading(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def add_inline_runs(paragraph, text):
    """Handles **bold** and `code` inline spans."""
    pattern = re.compile(r"(\*\*.+?\*\*|`.+?`)")
    parts = pattern.split(text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(9.5)
        else:
            paragraph.add_run(part)


def add_title_page():
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(140)
    run = p.add_run("FACTORY SAFETY SENTINEL")
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = ORANGE

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run("Final Project Report")
    run2.font.size = Pt(30)
    run2.font.bold = True
    run2.font.color.rgb = NAVY

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p3.paragraph_format.space_before = Pt(18)
    for line in ["Smart-Facility Incident Detection System", "Assessment Submission — 2026-08-29", "Repository: github.com/M7mdknh/smart-detector"]:
        r = p3.add_run(line + "\n")
        r.font.size = Pt(13)
        r.font.color.rgb = GRAY

    doc.add_page_break()


def add_table(rows):
    if not rows:
        return
    n_cols = len(rows[0])
    table = doc.add_table(rows=0, cols=n_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for i, row in enumerate(rows):
        cells = table.add_row().cells
        for j, cell_text in enumerate(row):
            cells[j].width = Inches(6.0 / n_cols)
            para = cells[j].paragraphs[0]
            para.paragraph_format.space_after = Pt(2)
            add_inline_runs(para, cell_text.strip())
            for run in para.runs:
                run.font.size = Pt(8.5)
            if i == 0:
                set_cell_shading(cells[j], "E2E8F0")
                for run in para.runs:
                    run.font.bold = True
    doc.add_paragraph()


def add_code_block(lines):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(6)
    para.paragraph_format.space_after = Pt(6)
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), "0F1C2E")
    pPr.append(shd)
    text = "\n".join(lines)
    run = para.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0xE2, 0xE8, 0xF0)


def parse_table_block(md_lines, start):
    rows = []
    i = start
    while i < len(md_lines) and md_lines[i].strip().startswith("|"):
        line = md_lines[i].strip()
        if re.match(r"^\|[\s:|-]+\|$", line):
            i += 1
            continue
        cells = [c.strip() for c in line.strip("|").split("|")]
        rows.append(cells)
        i += 1
    return rows, i


def main():
    add_title_page()
    md_lines = SRC.read_text().splitlines()
    i = 0
    first_h1_done = False
    while i < len(md_lines):
        line = md_lines[i]
        stripped = line.strip()

        if stripped.startswith("```"):
            code_lines = []
            i += 1
            while i < len(md_lines) and not md_lines[i].strip().startswith("```"):
                code_lines.append(md_lines[i])
                i += 1
            add_code_block(code_lines)
            i += 1
            continue

        if stripped.startswith("|"):
            rows, i = parse_table_block(md_lines, i)
            add_table(rows)
            continue

        if stripped.startswith("# "):
            if first_h1_done:
                doc.add_page_break()
            first_h1_done = True
            h = doc.add_heading(level=1)
            add_inline_runs(h, stripped[2:])
            for run in h.runs:
                run.font.color.rgb = NAVY
            i += 1
            continue

        if stripped.startswith("## "):
            h = doc.add_heading(level=1)
            add_inline_runs(h, stripped[3:])
            for run in h.runs:
                run.font.color.rgb = NAVY
                run.font.size = Pt(17)
            i += 1
            continue

        if stripped.startswith("### "):
            h = doc.add_heading(level=2)
            add_inline_runs(h, stripped[4:])
            for run in h.runs:
                run.font.color.rgb = ORANGE
                run.font.size = Pt(13)
            i += 1
            continue

        if stripped == "---":
            doc.add_paragraph().add_run("").add_break()
            i += 1
            continue

        img_match = re.match(r"^!\[.*?\]\((.+?)\)$", stripped)
        if img_match:
            img_path = (SRC.parent / img_match.group(1)).resolve()
            if img_path.exists():
                doc.add_picture(str(img_path), width=Inches(6.2))
                last_p = doc.paragraphs[-1]
                last_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            i += 1
            continue

        if stripped.startswith("- ") or stripped.startswith("* "):
            p = doc.add_paragraph(style="List Bullet")
            add_inline_runs(p, stripped[2:])
            i += 1
            continue

        if re.match(r"^\d+\.\s", stripped):
            p = doc.add_paragraph(style="List Number")
            add_inline_runs(p, re.sub(r"^\d+\.\s", "", stripped))
            i += 1
            continue

        if stripped == "":
            i += 1
            continue

        p = doc.add_paragraph()
        add_inline_runs(p, stripped)
        i += 1

    doc.save(str(OUT))
    print(f"Wrote {OUT} ({len(doc.paragraphs)} paragraphs, {len(doc.tables)} tables)")


if __name__ == "__main__":
    main()
